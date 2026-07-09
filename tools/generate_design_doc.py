from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "系统设计说明书_修订版.md"
DOCX_OUT = ROOT / "系统设计说明书_修订版.docx"
PDF_OUT = ROOT / "系统设计说明书_修订版.pdf"


def iter_blocks(text):
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith("```"):
            lang = line.strip("`").strip()
            code = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            yield ("code", lang, "\n".join(code))
            continue
        if line.lstrip().startswith("|"):
            table = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                table.append(lines[i])
                i += 1
            yield ("table", table)
            continue
        yield ("line", line)
        i += 1


def split_table(lines):
    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells):
            continue
        rows.append(cells)
    width = max(len(row) for row in rows) if rows else 0
    return [row + [""] * (width - len(row)) for row in rows]


def clean_inline(text):
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return text


def set_run_font(run, font_name="宋体", size=10.5, bold=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    set_run_font(run, size=9.5, bold=bold)


def build_docx(text):
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)

    for style_name, size in [("Heading 1", 18), ("Heading 2", 15), ("Heading 3", 12)]:
        style = styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)

    for block in iter_blocks(text):
        if block[0] == "line":
            line = block[1]
            if line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=0)
            elif line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=1)
            elif line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=2)
            elif line.startswith("> "):
                para = doc.add_paragraph()
                run = para.add_run(line[2:].strip())
                set_run_font(run, size=10)
                para.style = styles["Intense Quote"] if "Intense Quote" in styles else styles["Normal"]
            elif re.match(r"^\d+\.\s+", line):
                para = doc.add_paragraph(style="List Number")
                run = para.add_run(re.sub(r"^\d+\.\s+", "", line))
                set_run_font(run)
            elif line.startswith("- "):
                para = doc.add_paragraph(style="List Bullet")
                run = para.add_run(line[2:].strip())
                set_run_font(run)
            else:
                para = doc.add_paragraph()
                run = para.add_run(line)
                set_run_font(run)
        elif block[0] == "table":
            rows = split_table(block[1])
            if not rows:
                continue
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for r, row in enumerate(rows):
                for c, value in enumerate(row):
                    set_cell_text(table.cell(r, c), value, bold=(r == 0))
            doc.add_paragraph()
        elif block[0] == "code":
            para = doc.add_paragraph()
            for idx, line in enumerate(block[2].splitlines()):
                if idx:
                    para.add_run().add_break(WD_BREAK.LINE)
                run = para.add_run(line)
                set_run_font(run, font_name="等线", size=9)

    doc.save(DOCX_OUT)


def pdf_styles():
    font_path = Path(r"C:\Windows\Fonts\simsun.ttc")
    if not font_path.exists():
        font_path = Path(r"C:\Windows\Fonts\simhei.ttf")
    pdfmetrics.registerFont(TTFont("CN", str(font_path)))
    styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "CNTitle",
            parent=styles["Title"],
            fontName="CN",
            fontSize=20,
            leading=28,
            alignment=TA_CENTER,
            spaceAfter=10,
        ),
        "h1": ParagraphStyle(
            "CNH1",
            parent=styles["Heading1"],
            fontName="CN",
            fontSize=15,
            leading=22,
            spaceBefore=10,
            spaceAfter=6,
        ),
        "h2": ParagraphStyle(
            "CNH2",
            parent=styles["Heading2"],
            fontName="CN",
            fontSize=12.5,
            leading=18,
            spaceBefore=8,
            spaceAfter=4,
        ),
        "body": ParagraphStyle(
            "CNBody",
            parent=styles["BodyText"],
            fontName="CN",
            fontSize=9.5,
            leading=15,
            alignment=TA_LEFT,
            spaceAfter=4,
        ),
        "small": ParagraphStyle(
            "CNSmall",
            parent=styles["BodyText"],
            fontName="CN",
            fontSize=8,
            leading=11,
            spaceAfter=2,
        ),
        "code": ParagraphStyle(
            "CNCode",
            parent=styles["Code"],
            fontName="CN",
            fontSize=8.5,
            leading=12,
            leftIndent=6,
            backColor=colors.HexColor("#f5f5f5"),
            borderPadding=4,
            spaceAfter=6,
        ),
    }


def para(text, style):
    return Paragraph(clean_inline(text), style)


def build_pdf(text):
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=A4,
        leftMargin=16 * mm,
        rightMargin=16 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
    )
    story = []

    for block in iter_blocks(text):
        if block[0] == "line":
            line = block[1]
            if line.startswith("# "):
                story.append(para(line[2:].strip(), styles["title"]))
                story.append(Spacer(1, 4))
            elif line.startswith("## "):
                story.append(para(line[3:].strip(), styles["h1"]))
            elif line.startswith("### "):
                story.append(para(line[4:].strip(), styles["h2"]))
            elif line.startswith("> "):
                story.append(para(line[2:].strip(), styles["small"]))
            elif re.match(r"^\d+\.\s+", line):
                story.append(para(line, styles["body"]))
            elif line.startswith("- "):
                story.append(para("• " + line[2:].strip(), styles["body"]))
            else:
                story.append(para(line, styles["body"]))
        elif block[0] == "table":
            rows = split_table(block[1])
            if not rows:
                continue
            page_width = A4[0] - 32 * mm
            col_width = page_width / len(rows[0])
            data = [[para(cell, styles["small"]) for cell in row] for row in rows]
            table = Table(data, colWidths=[col_width] * len(rows[0]), repeatRows=1)
            table.setStyle(
                TableStyle(
                    [
                        ("FONTNAME", (0, 0), (-1, -1), "CN"),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e9eef7")),
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#9aa4b2")),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]
                )
            )
            story.append(table)
            story.append(Spacer(1, 5))
        elif block[0] == "code":
            code = block[2].replace("\n", "<br/>")
            story.append(Paragraph(clean_inline(code), styles["code"]))

    doc.build(story)


def main():
    text = SOURCE.read_text(encoding="utf-8")
    build_docx(text)
    build_pdf(text)
    print(DOCX_OUT)
    print(PDF_OUT)


if __name__ == "__main__":
    main()
